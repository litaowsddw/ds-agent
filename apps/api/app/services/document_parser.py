"""文档解析服务。

上传知识库文件时，API 层会先把二进制文件解析成纯文本，再交给 KnowledgeStore 做切片、
embedding 和向量索引。解析逻辑集中在这里，避免路由层出现大量文件格式分支。
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


@dataclass(slots=True)
class ParsedDocument:
    """解析后的文档内容。"""

    # title 是展示给用户的文档标题，默认取文件名。
    title: str

    # content 是抽取出的纯文本内容。
    content: str

    # source_format 是文件后缀或来源类型。
    source_format: str


class DocumentParser:
    """把常见文档格式解析为纯文本。"""

    def parse(self, filename: str, payload: bytes) -> ParsedDocument:
        """解析上传文件。"""

        if not payload:
            raise ValueError("上传文件不能为空")

        suffix = Path(filename).suffix.lower()
        title = Path(filename).stem or filename

        if suffix in {"", ".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
            content = self._decode_text(payload)
        elif suffix == ".pdf":
            content = self._parse_pdf(payload)
        elif suffix == ".docx":
            content = self._parse_docx(payload)
        else:
            raise ValueError(f"暂不支持的文件类型：{suffix or 'unknown'}")

        content = content.strip()
        if not content:
            raise ValueError("文件未解析出有效文本")

        return ParsedDocument(title=title, content=content, source_format=suffix or "text")

    def _decode_text(self, payload: bytes) -> str:
        """按常见编码解析文本文件。"""

        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return payload.decode(encoding)
            except UnicodeDecodeError:
                continue
        return payload.decode("utf-8", errors="ignore")

    def _parse_pdf(self, payload: bytes) -> str:
        """解析 PDF 文本。"""

        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("PDF 解析依赖 pypdf 未安装") from exc

        reader = PdfReader(BytesIO(payload))
        page_texts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(text for text in page_texts if text.strip())

    def _parse_docx(self, payload: bytes) -> str:
        """解析 DOCX 段落和表格文本。"""

        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("DOCX 解析依赖 python-docx 未安装") from exc

        document = Document(BytesIO(payload))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)


document_parser = DocumentParser()
